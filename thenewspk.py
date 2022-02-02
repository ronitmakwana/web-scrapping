import requests
from bs4 import BeautifulSoup
import os
import shutil
from docx import Document
from PIL import Image
import uuid
from docx.shared import Inches,RGBColor,Mm,Pt
from subprocess import Popen


document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(18)
# font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)


os.mkdir('images')
os.mkdir('im')
os.mkdir('kg')


url = 'https://www.thenews.com.pk/latest/category/national'

r = requests.get(url)

html_content = r.content

soup = BeautifulSoup(html_content,'html.parser')

links = soup.find('div',class_='detail-center')

a = links.find_all('h2')
for h2 in a:
    l = h2.find_all('a')
    for i in l:
        href = i.get('href')
        # print(href)
        r2 = requests.get(href)
        html_content2 = r2.content
        soup2 = BeautifulSoup(html_content2,'html.parser')
        # print(soup2.prettify)
        title = soup2.find('h1')
        tit = title.text
        p = document.add_paragraph()
        r = p.add_run()
        r.bold = True
        r.font.size = Pt(26)
        r.underline = True
        
        r.font.color.rgb = RGBColor(133, 33, 23)
    
        r.add_text('Title : '+tit)
        print(tit)
        date = soup2.find('div',class_='category-date')
        da = date.text
        p = document.add_paragraph()
        r = p.add_run()
        r.bold = True
        r.font.size = Pt(26)
        r.underline = True
        
        r.font.color.rgb = RGBColor(133, 33, 23)
    
        r.add_text('Date : '+da)
        print(da)
        img =soup2.find('div',class_='medium-insert-images norender-embeds ui-sortable')
        if not img:
            print('----- no image -----')
        else:
            src = img.select('img')
            
        
            newsrc = images = src[0]
            image = newsrc.attrs['src']
            print(image)
            im = requests.get(image,stream=True).content
        filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
        with open(filename,"wb+") as ik:
            ik.write(im)
            ff = Image.open(ik).convert('P').save('im/'+str(uuid.uuid4())+'.png')
            source = 'im/'
            destination = 'kg/'
    
            allfiles = os.listdir(source)
    
            for f in allfiles:
                shutil.move(source + f, destination + f)
                p = document.add_paragraph()
                r = p.add_run()
                r.add_picture(destination + f,width=Inches(12),height=Inches(4))
                # document.save('page2.docx')
                # document.add_page_break()

                print('-------')
        # content = soup2.find('div',class_='story-detail')
        p = soup2.find_all('p')
        for te in p:
            cont = te.text
            p = document.add_paragraph()
            r = p.add_run()
    
            r.font.size = Pt(18)    
            r.add_text(cont)
            document.save('idrw/pk.docx')
            # print(cont)



LIBRE_OFFICE = '/opt/libreoffice7.1/program/soffice'

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()


sample_doc = 'idrw/pk.docx'
out_folder = 'idrw'
convert_to_pdf(sample_doc, out_folder)

shutil.rmtree('images/')
shutil.rmtree('im/')
shutil.rmtree('kg/')
print("pdf saved")
    