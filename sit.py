from bs4 import BeautifulSoup
import requests
from docx import Document
from PIL import Image
from docx.shared import Pt,Mm,RGBColor,Inches
from subprocess import Popen
import os
import uuid
import shutil

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(18)
# font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)

os.mkdir('images')
os.mkdir('im')
os.mkdir('kg')

url = "https://economictimes.indiatimes.com/"

r = requests.get(url)

html_content = r.content

soup = BeautifulSoup(html_content,'html.parser')
# print(soup.prettify)

links = soup.find('div',class_='tabsContent clearfix')
ul = links.find('ul')
li = ul.find_all('a')[0:-1]
# print(li)
# for u in ul:
    # print(u.get_text())
# a = li.find_all('a')
for i in li:
    href = i.get('href')
    # print(href)
    url2 = 'https://economictimes.indiatimes.com/'+str(href)
    r2 = requests.get(url2)
    html_content2 = r2.content
    # print(html_content2)
    soup2 = BeautifulSoup(html_content2,'html.parser')
    # print(soup2.prettify)
    title = soup2.find('h1')
    tit = title.text
    print(tit)
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.font.size = Pt(26)
    r.underline = True
    
    r.font.color.rgb = RGBColor(133, 33, 23)
   
    r.add_text('Title : '+tit)
    date = soup2.find('time')
    if date == None:
        print('date')
        continue
    else:
        da = date.text
        p = document.add_paragraph()
        r = p.add_run()
        r.bold = True
        r.font.size = Pt(26)
        r.underline = True
        
        r.font.color.rgb = RGBColor(133, 33, 23)
    
        r.add_text('date : '+da)
        print(da)
    img =soup2.find('div',class_='imgBox')
    if not img:
        print('----- no image -----')
    else:
        src = img.select('img')
        
    
        newsrc = images = src[0]
        image = newsrc.attrs['src']
        im = requests.get(image,stream=True).content
    filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
    with open(filename,"wb+") as ik:
        ik.write(im)
        ff = Image.open(ik).convert('P').save('im/'+str(uuid.uuid4())+'.png')
        source = 'im/'
        destination = 'kg/'
  
        allfiles = os.listdir(source)
  
        for f in allfiles:
            shutil.move(source + f, destination + f)
            p = document.add_paragraph()
            r = p.add_run()
            r.add_picture(destination + f,width=Inches(12),height=Inches(4))
            # document.save('page2.docx')
            # document.add_page_break()

            print('-------')
    # if not img:
    #     print('----')
    # else:

        # print(image)
        content = soup2.find('div',class_='pageContent flt')
        if content == None:
            print('no content')
            continue
        else:
            con = content.text
        print(con)
        p = document.add_paragraph()
        r = p.add_run()
        # r.bold = True
        r.font.size = Pt(18)
        # r.underline = True
        
    #     # r.font.color.rgb = RGBColor(133, 33, 23)
    
        r.add_text(con)
        document.save('idrw/timesnew.docx')
        print('--------')
   
    
LIBRE_OFFICE = '/opt/libreoffice7.1/program/soffice'

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()


sample_doc = 'idrw/timesnew.docx'
out_folder = 'idrw'
convert_to_pdf(sample_doc, out_folder)

shutil.rmtree('images/')
shutil.rmtree('im/')
shutil.rmtree('kg/')
print("pdf saved")
