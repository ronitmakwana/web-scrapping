from bs4 import BeautifulSoup
import requests
from PIL import Image
from docx import Document
from docx.shared import Pt,Mm,RGBColor,Inches
from subprocess import Popen
import shutil
import uuid
import os

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(18)
# font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)


os.mkdir('images')
os.mkdir('im')
os.mkdir('kg')


url = 'https://www.thedefensepost.com/latest/'

r = requests.get(url)

html_content = r.content

soup = BeautifulSoup(html_content,'html.parser')
# print(soup.prettify)

title = soup.find('div',class_='mag-box-container')
# # print(title)
# for titles in title:

a = title.find_all('a',class_='more-link')
# print(a.get('href'))
for links in a:
    href = links.get('href')
    # print(href)
    r2 = requests.get(href)
    html_content2 = r2.content
    soup2 = BeautifulSoup(html_content2,'html.parser')
    tit = soup2.find('h1',class_='post-title entry-title') 
    h1 = tit.get_text()
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.font.size = Pt(26)
    r.underline = True
    
    r.font.color.rgb = RGBColor(133, 33, 23)
   
    r.add_text('Title : '+h1)
    # print(h1)
    date = soup2.find('span',class_='date meta-item')
    dates = date.text
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.font.size = Pt(26)
    r.underline = True
    
    r.font.color.rgb = RGBColor(133, 33, 23)
   
    r.add_text('Date : '+dates)
    print(dates)
    img = soup2.find('div',class_='featured-area')
    src = img.select('img')
        
    newsrc = images = src[0]
    image = newsrc.attrs['src']
    print(image)
    im = requests.get(image,stream=True).content
    filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
    with open(filename,"wb+") as ik:
        ik.write(im)
        ff = Image.open(ik).convert('P').save('im/'+str(uuid.uuid4())+'.png')
        source = 'im/'
        destination = 'kg/'
  
        allfiles = os.listdir(source)
  
        for f in allfiles:
            shutil.move(source + f, destination + f)
            p = document.add_paragraph()
            r = p.add_run()
            r.add_picture(destination + f,width=Inches(12),height=Inches(4))
    content = soup2.find('div',class_='entry-content entry clearfix')
    conts = content.text
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text(conts)
    document.save('idrw/post.docx')


LIBRE_OFFICE = '/opt/libreoffice7.1/program/soffice'

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()


sample_doc = 'idrw/post.docx'
out_folder = 'idrw'
convert_to_pdf(sample_doc, out_folder)

shutil.rmtree('images/')
shutil.rmtree('im/')
shutil.rmtree('kg/')
print("pdf saved")
    