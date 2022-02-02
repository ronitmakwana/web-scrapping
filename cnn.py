from bs4 import BeautifulSoup
import requests
import os
import uuid
from PIL import Image
import shutil
from subprocess import Popen
from docx import Document
from docx.shared import Mm,Pt,Inches,RGBColor




document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(18)
# font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)

os.mkdir('images')
os.mkdir('im')
os.mkdir('kg')

url = 'https://www.bbc.com/news/world/asia/china'

r = requests.get(url)

html_content = r.content

soup =BeautifulSoup(html_content,'html.parser')

# print(soup.prettify)

links = soup.find('div',class_='b-pw-1280')
# print(links)
title = links.find('div',class_='no-mpu')
# print(title)
href = title.find_all('a')
for i in href:
    hr = i.get('href')
    # print(hr)
    r2 = requests.get('https://www.bbc.com/'+str(hr))
    html_content2 = r2.content
    soup2 = BeautifulSoup(html_content2,'html.parser')
    # print(soup2.prettify)
    h1 = soup2.find('h1')
    text1 = h1.text
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.font.size = Pt(26)
    r.underline = True
    
    r.font.color.rgb = RGBColor(133, 33, 23)
   
    r.add_text('Title : '+text1)
    time = soup2.find('time')
    date = time.text
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.font.size = Pt(26)
    r.underline = True
    
    r.font.color.rgb = RGBColor(133, 33, 23)
   
    r.add_text('Title : '+date)
    img =soup2.find('div',class_='ssrcss-ab5fd8-StyledFigureContainer e34k3c21')
    if not img:
        print('----- no image -----')
    else:
        src = img.select('img')
        
    
        newsrc = images = src[0]
        image = newsrc.attrs['src']
        # print(image)
        im = requests.get(image,stream=True).content
    filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
    with open(filename,"wb+") as ik:
        ik.write(im)
        ff = Image.open(ik).convert('P').save('im/'+str(uuid.uuid4())+'.png')
        source = 'im/'
        destination = 'kg/'
  
        allfiles = os.listdir(source)
  
        for f in allfiles:
            shutil.move(source + f, destination + f)
            p = document.add_paragraph()
            r = p.add_run()
            r.add_picture(destination + f,width=Inches(12),height=Inches(4))
            # document.save('page2.docx')
            # document.add_page_break()

            print('-------')




    cont = soup2.find_all('div',class_='ssrcss-uf6wea-RichTextComponentWrapper e1xue1i85')
    for c in cont:
        co = c.get_text()
        p = document.add_paragraph()
        r = p.add_run()

        r.font.size = Pt(18)
   
        r.add_text(co)
        print(co)
        document.save('idrw/bbc1.docx')


LIBRE_OFFICE = '/usr/lib/libreoffice/program/soffice'

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()


sample_doc = 'idrw/bbc1.docx'
out_folder = 'idrw'
convert_to_pdf(sample_doc, out_folder)

shutil.rmtree('images/')
shutil.rmtree('im/')
shutil.rmtree('kg/')
print("pdf saved")
