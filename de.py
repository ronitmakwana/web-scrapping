from bs4 import BeautifulSoup
import requests
from PIL import Image
from docx import Document
from docx.shared import Pt,Mm,Inches,RGBColor
from subprocess import Popen
import uuid
import os
import shutil

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(18)
# font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)

os.mkdir('images')
os.mkdir('im')
os.mkdir('kg')


url = "http://www.indiandefensenews.in/"


r = requests.get(url)
html_content = r.content

soup = BeautifulSoup(html_content,'html.parser')
title = soup.find('div',class_='date-posts')
# print(title.prettify)
h1 = title.find_all('h1')
for h in h1:
    a = h.find('a')
    href = a.get('href')
    # print(href)
    # print(h.get_text())

    r2  = requests.get(href)
    html_content2 = r2.content
    soup2 = BeautifulSoup(html_content2,'html.parser')
    title2 = soup2.find('h1')
    tit = title2.text
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.font.size = Pt(26)
    r.underline = True
    
    r.font.color.rgb = RGBColor(133, 33, 23)
   
    r.add_text('Title : '+tit)
    # print(tit.text)
    date = title2.find_next('div',class_='post-details')
    da = date.text
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.font.size = Pt(26)
    r.underline = True
    r.font.color.rgb = RGBColor(133, 33, 23)
    tit = title2.text
    r.add_text('Date : '+da)
    print(da)
    cont = soup2.find('div',id='adsense-target')

    src = cont.select('img')
        
    newsrc = images = src[0]
    image = newsrc.attrs['src']
    # print(image)
    im = requests.get(image,stream=True).content
    filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
    with open(filename,"wb+") as ik:
        ik.write(im)
        ff = Image.open(ik).convert('P').save('im/'+str(uuid.uuid4())+'.png')
        source = 'im/'
        destination = 'kg/'
  
        allfiles = os.listdir(source)
  
        for f in allfiles:
            shutil.move(source + f, destination + f)
            p = document.add_paragraph()
            r = p.add_run()
            r.add_picture(destination + f,width=Inches(12),height=Inches(4))
            # document.save('page2.docx')
            # document.add_page_break()

            print('-------')
    c = cont.text
    p = document.add_paragraph()
    r = p.add_run()
    r.font.size = Pt(18)

    r.add_text(c)
    
    p = document.add_paragraph()
    r = p.add_run()
    document.save('idrw/ids1.docx')
    print('====')
    # print(c)

LIBRE_OFFICE = '/opt/libreoffice7.1/program/soffice'

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()


sample_doc = 'idrw/ids1.docx'
out_folder = 'idrw'
convert_to_pdf(sample_doc, out_folder)

shutil.rmtree('images/')
shutil.rmtree('im/')
shutil.rmtree('kg/')
print("pdf saved")


