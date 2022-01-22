from bs4 import BeautifulSoup
import requests
import pdfkit
import os
import uuid
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
from PIL import Image
import io
import shutil
import base64
from docx.shared import Mm

document = Document()

section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)
# section.left_margin = Mm(25.4)
# section.right_margin = Mm(25.4)
# section.top_margin = Mm(25.4)
# section.bottom_margin = Mm(25.4)
# section.header_distance = Mm(12.7)
# section.footer_distance = Mm(12.7)



r = requests.get('https://idrw.org/page/2/')
os.mkdir('images')
os.mkdir('im')
os.mkdir('kg')
htmlcontent = r.content
# print(htmlcontent)

soup = BeautifulSoup(htmlcontent,"html.parser")
# print(soup.get_text().replace(' ',''))




title = soup.find('div',class_='art-layout-cell art-content')

link  = title.find_all('a')

# print(link)

head = title.find_all('h2')
for h in head:
    title = h.text
    # bold = '\33[1m'
    print(f'title:{title}')
    p = document.add_paragraph()
    r = p.add_run()
    r.add_text(title)
    # file = 'imafd.docx'
    # with open(file,"a") as f:
    #     f.write(f'title:{title}'+'\n'+'\n')

    
               
    for link in h.find_all('a'):
        li = link.get('href')
            # print(li)
        links = requests.get(li)
        htmllink = links.content
        soup2 = BeautifulSoup(htmllink,"html.parser")
            # con = soup2.find('div',class_='art-layout-cell art-content')

        con = soup2.find('div',class_='art-content-layout')
        img = con.find('figure',class_='aligncenter size-large is-resized')
        src = img.select('img')
        
        k = []
        # for ls in src:
        #     k.append(ls)
        # print(k)
        # # k.pop(0)
        # print(k)

        newsrc = images = src[0]
        image = newsrc.attrs['src']
        # print(image)
        im = requests.get(image,stream=True).content
        


        # image = io.BytesIO(im.content)
        # im = im.decode('utf8')
        # r = requests.get(image, allow_redirects=True)
        
        # open('images/result_'+str(uuid.uuid4())+'.jpg', 'wb').write(r.content)
        # r.add_picture(r.content)
        # document.save('drws7.docx')
        # document.add_picture(open('1', mode='rb'))




        
        filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
        with open(filename,"wb+") as ik:
            ik.write(im)
            ff = Image.open(ik).convert('RGB').save('im/'+str(uuid.uuid4())+'.jpg')
            source = 'im/'
            destination = 'kg/'
  
            allfiles = os.listdir(source)
  
            for f in allfiles:
                shutil.move(source + f, destination + f)
                p = document.add_paragraph()
                r = p.add_run()
                r.add_picture(destination + f,width=Inches(12),height=Inches(4))
                # document.save('page2.docx')
                document.add_page_break()
                print('-------')


            # fina = 'kg/result_'+str(uuid.uuid4())+'.jpg'
            # with open(fina,"wb+") as gg:
            #     gg.write(fil)
            #     document.add_picture(gg)
           
            
            # image1 = Image.open(ik)
            # im1 = image1.convert('RGB')
            # im1.save('im/'+str(uuid.uuid4())+'.jpg') 
            # aa =[]
            
            # aa.append(im1)
            # print(aa)
            # for ii in aa:
            #     print(ii)
            #     r.add_picture(ii)
            # with open (r.add_picture(kk),'a') as bb:
            #     bb.write(kk)
        
               
            # document.save('idrws7a.docx')
        # print("----")

       
            
           
            s = con.find_all_next('p')
            span = con.find('span')
            t = span.text
            print(f'Date: {t}')
            p = document.add_paragraph()
            r = p.add_run()
            r.add_text(t)
                # f.write(f'Date: {t}'+'\n')
                
            for k in s:
                content = k.get_text()
                # print(content)
                p = document.add_paragraph()
                r = p.add_run()
                r.add_text(content)
              
                p = document.add_paragraph()
                r = p.add_run()
                document.save('page34.docx')

                
shutil.rmtree('images/')
shutil.rmtree('im/')
shutil.rmtree('kg/')

# convert('image.docx','cov.pdf') 
#                 f.write(content+'\n')
                
            

    


#     # pdf.add_page()
#     # pdf.set_font("Arial",size=12)
#     # pdf.cell(200,10, txt=titles,ln=1,align='C') 
#     # pdf.output("title.pdf")
# print('save')
# # def save_pdf(htmls, file_name):
    
# htmls = file
# file_name = "iaf.pdf"
# options = {

#     'page-size': 'Letter',

#     'margin-top': '0.075in',

#     'margin-right': '0.75in',

#     'margin-bottom': '0.75in',

#     'margin-left': '0.75in',

#     'encoding': "UTF-8",

#     'custom-header': [

#     ('Accept-Encoding', 'gzip')

# ],

#     'cookie': [

#         ('cookie-name1', 'cookie-value1'),

#         ('cookie-name2', 'cookie-value2'),

#         ],

#         'outline-depth': 10,

#         }

# pdfkit.from_file(htmls,file_name, options=options) 
print("pdf saved")
# os.remove(file)

    # print(h.text.split())

# for l in link:  

#     print(l.get('href'))
# a = title.h2.a.text
# print(a)

# for index,item in enumerate(title):
#     content = item.find_all('a')
#     for link in content:
#         print(link.get('href'))
   

    
    # for j in content:
    #     c = j.find('p')
    #     print(f'content:{c}')
# a = title.find('a')
# print(a)
# print(soup.prettify)
# print(a)
# for i in a:
#     ti = i.find('p')
#     print(ti)



# for i in title:
#     i = i.find('title')
#     print(i)




