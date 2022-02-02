from bs4 import BeautifulSoup
import requests
from docx import Document
from PIL import Image
from docx.shared import Pt,Mm,RGBColor,Inches
from subprocess import Popen
import os
import uuid
import shutil

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(18)
# font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)
os.mkdir('images')
os.mkdir('im')
os.mkdir('kg')
url = 'https://www.bhaskarhindi.com/state/jammu-and-kashmir'

r = requests.get(url)

html_content = r.content

soup = BeautifulSoup(html_content,'html.parser')

# print(soup)

links = soup.find('ul',class_='col-md-12 list')
# print(links)
a = links.find_all('a')
for link in a:
    href = link.get('href')
    print(href)

    url2 = ('https://www.bhaskarhindi.com/'+str(href))
    r2 = requests.get(url2)
    html_content2 = r2.content
    soup2 = BeautifulSoup(html_content2,'html.parser')
    # print(soup2)
    title = soup2.find('h1')
    if not title:
        print('no title')
    else:
       t = title.text
       p = document.add_paragraph()
       r = p.add_run()
       r.bold = True
       r.font.size = Pt(26)
       r.underline = True
	    
       r.font.color.rgb = RGBColor(133, 33, 23)
   
       r.add_text('Title : '+t)
       print(t)
    date = soup2.find('div',class_='time')
    if not date:

        print('no date')
    else:
        
       dt =date.text
       p = document.add_paragraph()
       r = p.add_run()
       r.bold = True
       r.font.size = Pt(26)
       r.underline = True
    
       r.font.color.rgb = RGBColor(133, 33, 23)
   
       r.add_text('Date : '+dt)
       print(dt)
    img =soup2.find('div',class_='featured')
    if not img:
        print('----- no image -----')
    else:
        src = img.select('img')
        
    
        newsrc = images = src[0]
        image = newsrc.attrs['src']
        print(image)
        im = requests.get(image,stream=True).content
        
    filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
    with open(filename,"wb+") as ik:
         ik.write(im)
         ff = Image.open(ik).convert('P').save('im/'+str(uuid.uuid4())+'.png')
         source = 'im/'
         destination = 'kg/'
  
         allfiles = os.listdir(source)
  
         for f in allfiles:
             shutil.move(source + f, destination + f)
             p = document.add_paragraph()
             r = p.add_run()
             r.add_picture(destination + f,width=Inches(12),height=Inches(4))
             # document.save('page2.docx')
             # document.add_page_break()

             print('-------')

   
    con = soup2.find('div',class_='description')
    if not con:
        print('no content')
    else:
        cont = con.text
        p = document.add_paragraph()
        r = p.add_run()
        r.font.size = Pt(18)
       
        r.add_text(cont)
        document.save('idrw/bhasker2.docx')
        print('----')

LIBRE_OFFICE = '/usr/lib/soffice'

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()


sample_doc = 'idrw/bhasker2.docx'
out_folder = 'idrw'
convert_to_pdf(sample_doc, out_folder)

shutil.rmtree('images/')
shutil.rmtree('im/')
shutil.rmtree('kg/')
	 
		

