from bs4 import BeautifulSoup
import requests
import os
import uuid
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
from PIL import Image
import shutil
from docx.shared import Mm,Pt
from subprocess import  Popen

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(20)
section = document.sections[0]
section.page_height = Mm(350)
section.page_width = Mm(400)




r = requests.get('https://idrw.org/page/2/')
os.mkdir('images')
os.mkdir('im')
os.mkdir('kg')
# os.mkdir('idrw')
htmlcontent = r.content
# print(htmlcontent)

soup = BeautifulSoup(htmlcontent,"html.parser")
# print(soup.get_text().replace(' ',''))




title = soup.find('div',class_='art-layout-cell art-content')

link  = title.find_all('a')

# print(link)

head = title.find_all('h2')
for h in head:
    title = h.text
    print(f'title:{title}')
    p = document.add_paragraph()
    r = p.add_run()
    r.bold = True
    r.underline=True
    r.add_text('Title :  '+title)
 
    
               
    for link in h.find_all('a'):
        li = link.get('href')
            # print(li)
        links = requests.get(li)
        htmllink = links.content
        soup2 = BeautifulSoup(htmllink,"html.parser")
            # con = soup2.find('div',class_='art-layout-cell art-content')

        con = soup2.find('div',class_='art-content-layout')
        img = con.find('figure',class_='aligncenter size-large is-resized')
        src = img.select('img')
        
        newsrc = images = src[0]
        image = newsrc.attrs['src']
        # print(image)
        im = requests.get(image,stream=True).content
        
        filename = 'images/result_'+str(uuid.uuid4())+'.jpg'
        with open(filename,"wb+") as ik:
            ik.write(im)
            ff = Image.open(ik).convert('P').save('im/'+str(uuid.uuid4())+'.png')
            source = 'im/'
            destination = 'kg/'
  
            allfiles = os.listdir(source)
  
            for f in allfiles:
                shutil.move(source + f, destination + f)
                p = document.add_paragraph()
                r = p.add_run()
                r.add_picture(destination + f,width=Inches(12),height=Inches(4))
                # document.save('page2.docx')
                document.add_page_break()
                print('-------')

            s = con.find_all_next('p')
            span = con.find('span')
            t = span.text
            print(f'Date: {t}')
           
            p = document.add_paragraph()
            r = p.add_run()
            r.underline = True
            r.bold = True
            r.add_text('Date : ' +t)
                
                
            for k in s:
                content = k.get_text()
                # print(content)
               
                p = document.add_paragraph()
                r = p.add_run()
                r.add_text(content)
              
                p = document.add_paragraph()
                r = p.add_run()
                document.save('idrw/11.docx')

### for linux only if microsoft office is not installed 

LIBRE_OFFICE = '/opt/libreoffice7.1/program/soffice'

def convert_to_pdf(input_docx, out_folder):
    p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
               out_folder, input_docx])
    print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
    p.communicate()


sample_doc = 'idrw/11.docx'
out_folder = 'idrw'
convert_to_pdf(sample_doc, out_folder)

shutil.rmtree('images/')
shutil.rmtree('im/')
shutil.rmtree('kg/')
print("pdf saved")




